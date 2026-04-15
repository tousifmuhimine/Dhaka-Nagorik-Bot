"""Generate complaint documents for download and email delivery."""

from __future__ import annotations

from pathlib import Path
from typing import Any, cast

from django.conf import settings
from django.utils import timezone
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.styles.style import ParagraphStyle as DocxParagraphStyle
from docx.text.paragraph import Paragraph as DocxParagraph
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


class ComplaintDocumentService:
    """Generate formal DOCX and PDF complaint applications."""

    def __init__(self):
        self.output_dir = Path(settings.DOCUMENT_OUTPUT_DIR)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(self, complaint, extracted_complaint=None, attachments=None) -> dict:
        """Generate both document formats and return their filesystem paths."""
        attachments = attachments or []
        timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
        base_name = f"complaint_{complaint.id}_{timestamp}"

        docx_path = self.output_dir / f"{base_name}.docx"
        pdf_path = self.output_dir / f"{base_name}.pdf"

        try:
            self._generate_docx(docx_path, complaint, extracted_complaint, attachments)
        except Exception as e:
            raise RuntimeError(f"Failed to generate DOCX: {str(e)}")
        
        try:
            self._generate_pdf(pdf_path, complaint, extracted_complaint, attachments)
        except Exception as e:
            raise RuntimeError(f"Failed to generate PDF: {str(e)}")

        return {
            'docx_path': str(docx_path),
            'pdf_path': str(pdf_path),
        }

    def _generate_docx(self, path: Path, complaint: Any, extracted_complaint: Any, attachments: list[Any]) -> None:
        """Create a formal Word application for the complaint."""
        document: Any = Document()
        self._configure_docx_page(document)

        heading = document.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading_run = heading.add_run('APPLICATION FOR CIVIC COMPLAINT RESOLUTION')
        heading_run.bold = True
        heading_run.font.size = Pt(15)

        ref_line = document.add_paragraph()
        ref_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ref_run = ref_line.add_run(f"Complaint Reference: DN-{complaint.id:05d}")
        ref_run.italic = True
        ref_run.font.size = Pt(10.5)

        document.add_paragraph()

        for line in self._recipient_lines(complaint):
            self._add_docx_paragraph(document, line)

        document.add_paragraph()
        self._add_docx_paragraph(document, f"Date: {self._format_display_date(timezone.now())}")
        document.add_paragraph()

        subject = self._build_subject_line(complaint)
        subject_paragraph = self._add_docx_paragraph(document, '')
        subject_run = subject_paragraph.add_run('Subject: ')
        subject_run.bold = True
        subject_paragraph.add_run(subject)

        document.add_paragraph()
        self._add_docx_paragraph(document, 'Sir/Madam,')
        document.add_paragraph()

        for paragraph in self._build_application_paragraphs(complaint, extracted_complaint, attachments):
            self._add_docx_paragraph(document, paragraph, justified=True)
            document.add_paragraph()

        self._add_docx_paragraph(document, 'Therefore, I respectfully request your office to inspect the location and take the necessary corrective action at the earliest possible time.', justified=True)
        document.add_paragraph()
        self._add_docx_paragraph(document, 'Sincerely,')
        document.add_paragraph()
        self._add_docx_paragraph(document, complaint.citizen.get_full_name() or complaint.citizen.username)
        self._add_docx_paragraph(document, f"Email: {complaint.citizen.email or 'Not provided'}")
        self._add_docx_paragraph(document, f"Submitted via Dhaka Nagorik AI on {complaint.created_at.strftime('%B %d, %Y %I:%M %p')}")

        self._add_docx_section_heading(document, 'Annexure A: Complaint Details')
        details_table = document.add_table(rows=0, cols=2)
        details_table.style = 'Table Grid'
        for label, value in self._build_detail_rows(complaint, extracted_complaint):
            row = details_table.add_row().cells
            row[0].text = label
            row[1].text = value or '-'

        if attachments:
            self._add_docx_section_heading(document, 'Annexure B: Attached Evidence')
            for index, attachment in enumerate(attachments, start=1):
                self._add_docx_paragraph(document, f"{index}. {attachment.original_name}")

        validation_notes = self._validation_notes(extracted_complaint)
        if validation_notes:
            self._add_docx_section_heading(document, 'Annexure C: Supporting Notes')
            for note in validation_notes:
                self._add_docx_paragraph(document, f"- {note}")

        document.save(str(path))

    def _generate_pdf(self, path: Path, complaint: Any, extracted_complaint: Any, attachments: list[Any]) -> None:
        """Create a formal PDF application for the complaint."""
        doc = SimpleDocTemplate(
            str(path),
            pagesize=A4,
            leftMargin=0.85 * inch,
            rightMargin=0.85 * inch,
            topMargin=0.85 * inch,
            bottomMargin=0.85 * inch,
        )
        styles = self._build_pdf_styles()
        story = []

        story.append(Paragraph('APPLICATION FOR CIVIC COMPLAINT RESOLUTION', styles['title']))
        story.append(Paragraph(f"Complaint Reference: DN-{complaint.id:05d}", styles['reference']))
        story.append(Spacer(1, 10))

        for line in self._recipient_lines(complaint):
            story.append(Paragraph(self._escape_pdf(line), styles['body']))
        story.append(Spacer(1, 10))
        story.append(Paragraph(self._escape_pdf(f"Date: {self._format_display_date(timezone.now())}"), styles['body']))
        story.append(Spacer(1, 10))
        story.append(Paragraph(
            f"<b>Subject:</b> {self._escape_pdf(self._build_subject_line(complaint))}",
            styles['body'],
        ))
        story.append(Spacer(1, 10))
        story.append(Paragraph('Sir/Madam,', styles['body']))
        story.append(Spacer(1, 10))

        for paragraph in self._build_application_paragraphs(complaint, extracted_complaint, attachments):
            story.append(Paragraph(self._escape_pdf(paragraph), styles['justified']))
            story.append(Spacer(1, 8))

        request_text = 'Therefore, I respectfully request your office to inspect the location and take the necessary corrective action at the earliest possible time.'
        story.append(Paragraph(self._escape_pdf(request_text), styles['justified']))
        story.append(Spacer(1, 16))
        story.append(Paragraph('Sincerely,', styles['body']))
        story.append(Spacer(1, 16))
        story.append(Paragraph(self._escape_pdf(complaint.citizen.get_full_name() or complaint.citizen.username), styles['body']))
        story.append(Paragraph(self._escape_pdf(f"Email: {complaint.citizen.email or 'Not provided'}"), styles['body']))
        story.append(Paragraph(self._escape_pdf(f"Submitted via Dhaka Nagorik AI on {complaint.created_at.strftime('%B %d, %Y %I:%M %p')}"), styles['body']))
        story.append(Spacer(1, 18))

        story.append(Paragraph('Annexure A: Complaint Details', styles['section']))
        details_table = Table(
            [[Paragraph(f"<b>{self._escape_pdf(label)}</b>", styles['table_label']), Paragraph(self._escape_pdf(value or '-'), styles['table_value'])]
             for label, value in self._build_detail_rows(complaint, extracted_complaint)],
            colWidths=[1.8 * inch, 4.6 * inch],
            hAlign='LEFT',
        )
        details_table.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.6, colors.HexColor('#B6BBC6')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F4F6F9')),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        story.append(details_table)

        if attachments:
            story.append(Spacer(1, 16))
            story.append(Paragraph('Annexure B: Attached Evidence', styles['section']))
            for index, attachment in enumerate(attachments, start=1):
                story.append(Paragraph(self._escape_pdf(f"{index}. {attachment.original_name}"), styles['body']))

        validation_notes = self._validation_notes(extracted_complaint)
        if validation_notes:
            story.append(Spacer(1, 16))
            story.append(Paragraph('Annexure C: Supporting Notes', styles['section']))
            for note in validation_notes:
                story.append(Paragraph(self._escape_pdf(f"- {note}"), styles['body']))

        doc.build(story)

    def _configure_docx_page(self, document: Any) -> None:
        """Set page margins and default font for generated DOCX files."""
        section = document.sections[0]
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

        normal_style = cast(DocxParagraphStyle, document.styles['Normal'])
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(11)

    def _build_subject_line(self, complaint: Any) -> str:
        """Create a formal subject line for the application."""
        return (
            f"Prayer for urgent action regarding {complaint.get_category_display().lower()} issue at "
            f"{complaint.area}, {complaint.thana}"
        )

    def _recipient_lines(self, complaint: Any) -> list[str]:
        """Build the recipient block for the authority application."""
        if complaint.assigned_authority:
            profile = getattr(complaint.assigned_authority, 'userprofile', None)
            lines = [
                'To',
                complaint.assigned_authority.get_full_name() or complaint.assigned_authority.username,
            ]
            if profile and profile.department:
                lines.append(profile.department)
            lines.append(f"Authority responsible for {complaint.thana}")
            lines.append('Dhaka Nagorik AI Complaint Desk')
            return lines

        return [
            'To',
            'Concerned Authority',
            f'Office responsible for {complaint.thana}',
            'Dhaka Nagorik AI Complaint Desk',
        ]

    def _build_application_paragraphs(self, complaint: Any, extracted_complaint: Any, attachments: list[Any]) -> list[str]:
        """Compose the main body paragraphs of the application."""
        citizen_name = complaint.citizen.get_full_name() or complaint.citizen.username
        duration = ''
        if extracted_complaint and extracted_complaint.duration:
            duration = extracted_complaint.duration.strip()

        opening = (
            f"I, {citizen_name}, am submitting this application to formally report a "
            f"{complaint.get_category_display().lower()}-related civic problem located at {complaint.area}, {complaint.thana}."
        )
        if duration:
            opening += f" According to the submitted information, the problem has persisted for {duration}."

        body = complaint.description.strip()
        if extracted_complaint and extracted_complaint.full_description:
            summary = extracted_complaint.full_description.strip()
        else:
            summary = ''

        impact = (
            "The matter is affecting public convenience, safety, and normal movement in the area, "
            "and it requires prompt review by the responsible office."
        )

        evidence = ''
        if attachments:
            evidence = (
                f"For supporting verification, {len(attachments)} photo evidence file"
                f"{'s have' if len(attachments) != 1 else ' has'} been attached with this application."
            )

        paragraphs = [opening, body]
        if summary and summary.casefold() != body.casefold():
            paragraphs.append(f"Additional summary of the issue: {summary}")
        paragraphs.append(impact)
        if evidence:
            paragraphs.append(evidence)
        return paragraphs

    def _build_detail_rows(self, complaint: Any, extracted_complaint: Any) -> list[tuple[str, str]]:
        """Build the annexure rows shown below the formal letter."""
        rows = [
            ('Complaint ID', f'DN-{complaint.id:05d}'),
            ('Submitted By', complaint.citizen.get_full_name() or complaint.citizen.username),
            ('Citizen Email', complaint.citizen.email or 'Not provided'),
            ('Category', complaint.get_category_display()),
            ('Thana', complaint.thana),
            ('Area / Location', complaint.area),
            ('Current Status', complaint.get_status_display()),
            ('Submission Time', complaint.created_at.strftime('%B %d, %Y %I:%M %p')),
        ]
        if complaint.assigned_authority:
            rows.append(('Assigned Authority', complaint.assigned_authority.get_full_name() or complaint.assigned_authority.username))
        if extracted_complaint:
            rows.extend([
                ('Issue Duration', extracted_complaint.duration or '-'),
                ('Keywords', ', '.join(extracted_complaint.keywords or []) or '-'),
                ('Policy Reference', extracted_complaint.policy_reference or '-'),
            ])
        return rows

    def _validation_notes(self, extracted_complaint: Any) -> list[str]:
        """Collect optional AI validation notes for the annexure."""
        if not extracted_complaint or not isinstance(extracted_complaint.web_search_results, dict):
            return []

        notes = []
        validation = extracted_complaint.web_search_results
        if validation.get('recommendation'):
            notes.append(validation['recommendation'])
        notes.extend(validation.get('inconsistencies', []))
        return [note for note in notes if note]

    def _add_docx_paragraph(self, document: Any, text: str, justified: bool = False) -> DocxParagraph:
        """Add a paragraph with consistent application formatting."""
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justified else WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.25
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        return paragraph

    def _add_docx_section_heading(self, document: Any, text: str) -> None:
        """Add a simple section heading for the annexures."""
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

    def _build_pdf_styles(self) -> dict[str, ParagraphStyle]:
        """Create consistent PDF styles for the formal application."""
        styles = getSampleStyleSheet()
        return {
            'title': ParagraphStyle(
                'ApplicationTitle',
                parent=styles['Title'],
                fontName='Times-Bold',
                fontSize=14,
                leading=18,
                alignment=1,
                spaceAfter=4,
            ),
            'reference': ParagraphStyle(
                'Reference',
                parent=styles['Normal'],
                fontName='Times-Italic',
                fontSize=10,
                leading=12,
                alignment=1,
                spaceAfter=8,
            ),
            'body': ParagraphStyle(
                'ApplicationBody',
                parent=styles['Normal'],
                fontName='Times-Roman',
                fontSize=11,
                leading=15,
                spaceAfter=0,
            ),
            'justified': ParagraphStyle(
                'ApplicationBodyJustified',
                parent=styles['Normal'],
                fontName='Times-Roman',
                fontSize=11,
                leading=15,
                alignment=4,
            ),
            'section': ParagraphStyle(
                'SectionHeading',
                parent=styles['Heading3'],
                fontName='Times-Bold',
                fontSize=11.5,
                leading=14,
                spaceBefore=4,
                spaceAfter=8,
            ),
            'table_label': ParagraphStyle(
                'TableLabel',
                parent=styles['Normal'],
                fontName='Times-Bold',
                fontSize=10.5,
                leading=13,
            ),
            'table_value': ParagraphStyle(
                'TableValue',
                parent=styles['Normal'],
                fontName='Times-Roman',
                fontSize=10.5,
                leading=13,
            ),
        }

    def _escape_pdf(self, text: str) -> str:
        """Escape simple XML characters for reportlab paragraphs."""
        return (
            (text or '')
            .replace('&', '&amp;')
            .replace('<', '&lt;')
            .replace('>', '&gt;')
        )

    def _format_display_date(self, value: Any) -> str:
        """Format dates consistently for the application."""
        return value.strftime('%B %d, %Y')
