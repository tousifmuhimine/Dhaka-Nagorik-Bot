from __future__ import annotations

from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from app.core.config import get_settings
from app.schemas.complaint import ComplaintRecord


class DocumentGenerator:
    def __init__(self) -> None:
        self.settings = get_settings()

    async def generate_submission_documents(self, complaint: ComplaintRecord) -> tuple[str, str]:
        pdf_path = self.settings.document_output_dir / f"{complaint.id}.pdf"
        docx_path = self.settings.document_output_dir / f"{complaint.id}.docx"
        self._build_pdf(pdf_path, complaint)
        self._build_docx(docx_path, complaint)
        return str(pdf_path), str(docx_path)

    @staticmethod
    def _build_pdf(path: Path, complaint: ComplaintRecord) -> None:
        c = canvas.Canvas(str(path), pagesize=A4)
        text = c.beginText(50, 800)
        text.setFont("Helvetica", 11)
        lines = [
            "Dhaka Nagorik Complaint Submission",
            f"Complaint ID: {complaint.id}",
            f"Category: {complaint.category}",
            f"Thana: {complaint.thana}",
            f"Area: {complaint.area}",
            f"Urgency: {complaint.urgency}",
            f"Status: {complaint.status}",
            f"Created At: {complaint.created_at.isoformat()}",
            "",
            "Summary:",
            complaint.summary,
            "",
            "Original Complaint:",
            complaint.original_text,
        ]
        for line in lines:
            text.textLine(line[:110])
        c.drawText(text)
        c.showPage()
        c.save()

    @staticmethod
    def _build_docx(path: Path, complaint: ComplaintRecord) -> None:
        doc = Document()
        doc.add_heading("Dhaka Nagorik Complaint Submission", level=1)
        doc.add_paragraph(f"Complaint ID: {complaint.id}")
        doc.add_paragraph(f"Category: {complaint.category}")
        doc.add_paragraph(f"Thana: {complaint.thana}")
        doc.add_paragraph(f"Area: {complaint.area}")
        doc.add_paragraph(f"Urgency: {complaint.urgency}")
        doc.add_paragraph(f"Status: {complaint.status}")
        doc.add_paragraph(f"Created At: {complaint.created_at.isoformat()}")
        doc.add_paragraph("Summary:")
        doc.add_paragraph(complaint.summary)
        doc.add_paragraph("Original Complaint:")
        doc.add_paragraph(complaint.original_text)
        doc.save(str(path))
